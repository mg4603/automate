from sys import exit
try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Cm
except ImportError:
    exit('This program requires python-docx to run.')


def hello_world():
    doc = Document()
    doc.add_paragraph('Hello, world!')
    doc.save('hello_world.docx')

def multiple_paragraphs():
    doc = Document()
    doc.add_paragraph('Hello, World!')
    para_obj1 = doc.add_paragraph('This is a second paragraph.')
    para_obj2 = doc.add_paragraph('This is yet another paragraph.')
    para_obj1.add_run(' This text is being added to the second paragraph.')
    doc.save('multiple_paragraphs.docx')

def headings():
    doc = Document()
    doc.add_heading('Header 0', 0)
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)
    doc.save('headings.docx')

def two_page():
    doc = Document()
    doc.add_paragraph('This is on the first page!')
    doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)
    doc.add_paragraph('This is on the second page!')
    doc.save('two_page.docx')

def img_docx():
    doc = Document()
    doc.add_picture('zophie.png', width=Cm(4), height=Cm(4))
    doc.save('img_doc.docx')

# img_docx()
# two_page()
# headings()
# multiple_paragraphs()
# hello_world()