from argparse import ArgumentParser
from sys import exit
from pathlib import Path
from logging import debug, DEBUG, CRITICAL, disable, basicConfig
try:
    from docx import Document
    from docx.enum.text import WD_BREAK
except ImportError:
    exit('This program requires python-docx module to run.')

basicConfig(level=DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
disable(CRITICAL)

def parse_args():
    parser = ArgumentParser()
    parser.add_argument(
        'guest_list', metavar='guest-list',
        help='plaintext file with list of guests to create invitations for'
    )
    parser.add_argument(
        'invitation_doc', metavar='invitation-doc',
        help='word document of invitations to write to.'
    )
    return parser.parse_args()

def main():
    print('Custom Invitations')
    print()

    args = parse_args()
    guests_file_name = args.guest_list
    invitation_doc = args.invitation_doc
    
    guests_file_path = Path(guests_file_name)
    debug(invitation_doc)

    doc = Document(invitation_doc)
    with guests_file_path.open('r') as file:
        for guest in file.readlines():
            para_obj1 = doc.add_paragraph(
                'It would be a pleasure to have the company of'
            )
            para_obj1.style = 'non_user_content'
            para_obj1.alignment = 1

            para_obj2 = doc.add_paragraph()
            para_obj2_run = para_obj2.add_run(guest)
            para_obj2.style = 'user_content'
            para_obj2_run.bold = True
            para_obj2.alignment = 1

            

            para_obj3 = doc.add_paragraph(
                'at 11010 Memory Lane on the Evening of'
            )
            para_obj3.style = 'non_user_content'
            para_obj3.alignment = 1


            para_obj4 = doc.add_paragraph(
                'April 1st'
            )
            para_obj4.style = 'non_user_content'

            para_obj4.runs[0].add_break(WD_BREAK.PAGE)
            para_obj4.alignment = 1

    
    doc.save(invitation_doc)


if __name__ == '__main__':
    main()